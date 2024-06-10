import os
from tkinter import Tk, filedialog, simpledialog, messagebox
from docx import Document
from docx.shared import Inches
from PIL import Image

def select_folder():
    root = Tk()
    root.withdraw()  # Hide the root window
    folder_path = filedialog.askdirectory(title="Select Folder Containing Images")
    root.destroy()
    return folder_path

def get_image_width_cm():
    root = Tk()
    root.withdraw()  # Hide the root window
    width_cm = simpledialog.askfloat("Input", "Enter the width of the images in cm:", minvalue=1.0)
    root.destroy()
    return width_cm

def create_word_doc_with_images(folder_path, image_width_cm):
    # Create a new Document
    doc = Document()

    # Set the page size to A4
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)

    # Set narrow margins (0.5 inches on each side)
    margin_size = Inches(0.5)
    section.top_margin = margin_size
    section.bottom_margin = margin_size
    section.left_margin = margin_size
    section.right_margin = margin_size
    
    # Convert the width from cm to inches
    image_width_in = image_width_cm / 2.54
    
    # Calculate available width for images on a line
    available_width = section.page_width - section.left_margin - section.right_margin

    # Initialize a paragraph to add images in line
    paragraph = doc.add_paragraph()

    current_line_width = 0

    # Loop through the images in the folder
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff')):
            image_path = os.path.join(folder_path, filename)
            
            # Open the image to get the aspect ratio
            with Image.open(image_path) as img:
                width, height = img.size
                aspect_ratio = height / width
            
            # Check if adding this image would exceed the available width
            if current_line_width + image_width_in > available_width:
                # If it would, start a new paragraph
                paragraph = doc.add_paragraph()
                current_line_width = 0
            
            # Add image to the document with the specified width, maintaining the aspect ratio
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(image_width_in))
            
            # Update the current line width
            current_line_width += image_width_in

    # Save the document
    save_path = os.path.join(folder_path, "arranged_images.docx")
    doc.save(save_path)
    
    # Show success message
    messagebox.showinfo("Success", f"Document saved as {save_path}")
    
    # Open the created document
    os.startfile(save_path)

if __name__ == "__main__":
    folder_path = select_folder()
    if folder_path:
        image_width_cm = get_image_width_cm()
        if image_width_cm:
            create_word_doc_with_images(folder_path, image_width_cm)
        else:
            print("No width specified.")
    else:
        print("No folder selected.")


### wrap text to square

# import os
# from tkinter import Tk, filedialog, simpledialog, messagebox
# from docx import Document
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from docx.shared import Inches
# from PIL import Image

# def select_folder():
#     root = Tk()
#     root.withdraw()  # Hide the root window
#     folder_path = filedialog.askdirectory(title="Select Folder Containing Images")
#     root.destroy()
#     return folder_path

# def get_image_width_cm():
#     root = Tk()
#     root.withdraw()  # Hide the root window
#     width_cm = simpledialog.askfloat("Input", "Enter the width of the images in cm:", minvalue=1.0)
#     root.destroy()
#     return width_cm

# def set_wrap_square(picture):
#     # Access the inline element and convert it to an anchor
#     inline = picture._inline
#     anchor = OxmlElement('wp:anchor')
#     for key, value in inline.attrib.items():
#         anchor.set(key, value)
    
#     # Set the wrap type to square
#     wrap_square = OxmlElement('wp:wrapSquare')
#     wrap_square.set(qn('w:wrapText'), 'bothSides')
#     anchor.append(wrap_square)
    
#     # Move children from inline to anchor
#     for child in inline:
#         anchor.append(child)
    
#     # Replace the inline element with the anchor
#     inline.getparent().replace(inline, anchor)

# def create_word_doc_with_images(folder_path, image_width_cm):
#     # Create a new Document
#     doc = Document()

#     # Set the page size to A4
#     section = doc.sections[0]
#     section.page_height = Inches(11.69)
#     section.page_width = Inches(8.27)

#     # Set narrow margins (0.5 inches on each side)
#     margin_size = Inches(0.5)
#     section.top_margin = margin_size
#     section.bottom_margin = margin_size
#     section.left_margin = margin_size
#     section.right_margin = margin_size
    
#     # Convert the width from cm to inches
#     image_width_in = image_width_cm / 2.54

#     # Loop through the images in the folder
#     for filename in os.listdir(folder_path):
#         if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff')):
#             image_path = os.path.join(folder_path, filename)
            
#             # Open the image to get the aspect ratio
#             with Image.open(image_path) as img:
#                 width, height = img.size
#                 aspect_ratio = height / width
                
#             # Add image to the document with the specified width, maintaining the aspect ratio
#             picture = doc.add_picture(image_path, width=Inches(image_width_in))
#             set_wrap_square(picture)
            
#             # Add a paragraph break
#             doc.add_paragraph()

#     # Save the document
#     save_path = os.path.join(folder_path, "arranged_images.docx")
#     doc.save(save_path)
    
#     # Show success message
#     messagebox.showinfo("Success", f"Document saved as {save_path}")
    
#     # Open the created document
#     os.startfile(save_path)

# if __name__ == "__main__":
#     folder_path = select_folder()
#     if folder_path:
#         image_width_cm = get_image_width_cm()
#         if image_width_cm:
#             create_word_doc_with_images(folder_path, image_width_cm)
#         else:
#             print("No width specified.")
#     else:
#         print("No folder selected.")


###